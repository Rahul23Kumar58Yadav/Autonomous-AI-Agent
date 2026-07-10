
import os
import re
import logging
from datetime import datetime
from typing import Dict, Any, List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

logger = logging.getLogger("docgen.docx_builder")

OUTPUT_DIR = "outputs"

# Maps step output_keys to human-readable section headings.
# Anything not in this map falls back to a title-cased version of the key.
SECTION_TITLE_OVERRIDES = {
    "executive_summary": "Executive Summary",
    "assumptions_note": "Assumptions & Interpretation",
    "main_content": "Overview",
    "review_notes": "Review Notes",
}

ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)  # professional dark blue


class DocxBuilder:
    def __init__(self, output_dir: str = OUTPUT_DIR):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def build(
        self,
        document_type: str,
        title: str,
        assumptions: List[str],
        outputs: Dict[str, Any],
        step_log: List[Dict[str, Any]] = None,
    ) -> str:
        """
        Builds the .docx file and returns its filesystem path.
        """
        doc = Document()
        self._configure_base_styles(doc)

        self._add_title_page(doc, document_type, title)
        self._add_assumptions_section(doc, assumptions)

        for output_key, content in outputs.items():
            heading = self._resolve_heading(output_key)
            doc.add_heading(heading, level=1)
            self._render_content_block(doc, content)

        if step_log:
            self._add_appendix(doc, step_log)

        filepath = self._save(doc, title)
        logger.info(f"Document generated at {filepath}")
        return filepath

    # ------------------------------------------------------------------ #
    # Styling
    # ------------------------------------------------------------------ #

    def _configure_base_styles(self, doc: Document) -> None:
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

        # Heading 1 styling — accent color, slightly larger
        h1 = doc.styles["Heading 1"]
        h1.font.size = Pt(16)
        h1.font.color.rgb = ACCENT_COLOR
        h1.font.bold = True

    def _add_title_page(self, doc: Document, document_type: str, title: str) -> None:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = ACCENT_COLOR

        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle_para.add_run(document_type.title())
        sub_run.font.size = Pt(14)
        sub_run.font.italic = True

        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(
            f"Generated {datetime.now().strftime('%B %d, %Y %H:%M')} — Autonomous Agent"
        )
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        doc.add_page_break()

    def _add_assumptions_section(self, doc: Document, assumptions: List[str]) -> None:
        if not assumptions:
            return
        doc.add_heading("Assumptions Made by the Agent", level=1)
        intro = doc.add_paragraph(
            "Where the original request was ambiguous or incomplete, the "
            "agent made the following reasonable assumptions:"
        )
        intro.runs[0].font.italic = True
        for item in assumptions:
            doc.add_paragraph(item, style="List Bullet")

    def _add_appendix(self, doc: Document, step_log: List[Dict[str, Any]]) -> None:
        doc.add_page_break()
        doc.add_heading("Appendix: Agent Execution Log", level=1)
        doc.add_paragraph(
            "The following steps were autonomously planned and executed to "
            "produce this document.",
        )

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "Step"
        hdr[1].text = "Action"
        hdr[2].text = "Output"
        hdr[3].text = "Status"

        for step in step_log:
            row = table.add_row().cells
            row[0].text = str(step.get("id", ""))
            row[1].text = str(step.get("action", ""))
            row[2].text = str(step.get("output_key", ""))
            row[3].text = str(step.get("status", ""))

    # ------------------------------------------------------------------ #
    # Content rendering
    # ------------------------------------------------------------------ #

    def _render_content_block(self, doc: Document, content: Any) -> None:
        """
        Dispatches based on content shape:
        - list[dict] -> rendered as a table (from 'generate_table' steps)
        - str -> rendered as paragraphs, with '- ' lines as bullets
        """
        if isinstance(content, list) and content and isinstance(content[0], dict):
            self._render_table(doc, content)
        elif isinstance(content, str):
            self._render_text(doc, content)
        else:
            # Fallback for unexpected shapes — never silently drop content
            doc.add_paragraph(str(content))

    def _render_text(self, doc: Document, text: str) -> None:
        lines = text.strip().split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            elif re.match(r"^\d+\.\s", line):
                doc.add_paragraph(re.sub(r"^\d+\.\s", "", line), style="List Number")
            else:
                doc.add_paragraph(line)

    def _render_table(self, doc: Document, rows: List[Dict[str, Any]]) -> None:
        columns = list(rows[0].keys())
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr_cells[i].text = str(col).replace("_", " ").title()
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        for row_data in rows:
            row_cells = table.add_row().cells
            for i, col in enumerate(columns):
                row_cells[i].text = str(row_data.get(col, ""))

        doc.add_paragraph()  # spacing after table

    # ------------------------------------------------------------------ #
    # Helpers
    # ------------------------------------------------------------------ #

    def _resolve_heading(self, output_key: str) -> str:
        if output_key in SECTION_TITLE_OVERRIDES:
            return SECTION_TITLE_OVERRIDES[output_key]
        return output_key.replace("_", " ").title()

    def _save(self, doc: Document, title: str) -> str:
        safe_title = re.sub(r"[^a-zA-Z0-9_\- ]", "", title).strip().replace(" ", "_")
        safe_title = safe_title or "document"
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{safe_title}_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath